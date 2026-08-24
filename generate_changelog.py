#!/usr/bin/env python3
import subprocess
import re
import os
import sys

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("\nError: 'python-docx' is not installed.")
    print("Please run: pip install python-docx")
    print("Then execute this script again.\n")
    sys.exit(1)

def run_git_log():
    # Execute the live git command on the local repository
    try:
        result = subprocess.run(
            ["git", "log", "--reverse", "--oneline"],
            capture_output=True,
            text=True,
            check=True
        )
        return result.stdout.strip()
    except subprocess.CalledProcessError as e:
        print(f"\nError: Failed to execute git command. {e.stderr}")
        print("Make sure you are running this script inside your git repository.\n")
        sys.exit(1)
    except FileNotFoundError:
        print("\nError: 'git' command not found. Ensure Git is installed and in your PATH.\n")
        sys.exit(1)

def parse_commits(raw_log):
    commits = []
    for line in raw_log.split("\n"):
        line = line.strip()
        if not line:
            continue
        # Format: 'f003895 message' or 'f003895 (decorations) message'
        match = re.match(r"^([a-f0-9]{7,10})\s+(.*)$", line)
        if match:
            sha, msg = match.groups()
            # Strip branch/tag decorations like "(HEAD -> master, origin/master)"
            msg_clean = re.sub(r"^\([^\)]+\)\s*", "", msg).strip()
            commits.append((sha, msg_clean))
    return commits

def set_cell_background(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
        node = OxmlElement(m)
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_horizontal_border(paragraph, hex_color="A38955", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), hex_color)
    pBdr.append(bottom)
    pPr.append(pBdr)

def build_docx(commits, output_name="bbs-complete-commit-changelog.docx"):
    doc = Document()
    
    # Configure wide, clean margins (Elegant showroom style)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    COLOR_CHARCOAL = RGBColor(43, 43, 43)    # #2B2B2B
    COLOR_BRASS = RGBColor(163, 137, 85)     # #A38955
    HEX_IVORY = "F5F2EB"
    HEX_WHITE = "FFFFFF"
    HEX_CHARCOAL = "2B2B2B"
    
    # Header block
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("BENJAMIN BARKER STUDIOS")
    title_run.font.name = "Georgia"
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = COLOR_BRASS
    
    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_after = Pt(24)
    add_horizontal_border(sub_p, "A38955", "18")
    sub_run = sub_p.add_run("Complete Unified Repository History  ·  Chronological Changelog")
    sub_run.font.name = "Arial"
    sub_run.font.size = Pt(11)
    sub_run.font.italic = True
    sub_run.font.color.rgb = COLOR_CHARCOAL
    
    # Descriptive introduction
    intro_p = doc.add_paragraph()
    intro_p.paragraph_format.space_after = Pt(16)
    intro_run = intro_p.add_run(
        "A comprehensive, chronological repository ledger tracing all development cycles, interactive mechanics, "
        "and architectural integrations from the conception of the app to the most recent showroom releases."
    )
    intro_run.font.name = "Georgia"
    intro_run.font.size = Pt(11.5)
    intro_run.font.color.rgb = COLOR_CHARCOAL
    
    # Construct changelog table
    table = doc.add_table(rows=len(commits) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    
    col_widths = [Inches(1.2), Inches(5.3)]
    
    # Format table header row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["COMMIT", "DEVELOPMENT LOG & CONTRIBUTION DESCRIPTION"]
    for i, cell in enumerate(hdr_cells):
        cell.width = col_widths[i]
        set_cell_background(cell, HEX_CHARCOAL)
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 1 else WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(hdr_titles[i])
        run.font.name = "Georgia"
        run.font.size = Pt(9.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(245, 242, 235)
        
    # Prevent header from splitting
    trPr = table.rows[0]._tr.get_or_add_trPr()
    trPr.append(OxmlElement('w:tblHeader'))
    
    # Populate rows with alternating zebra striping
    for idx, (sha, msg) in enumerate(commits):
        row_cells = table.rows[idx + 1].cells
        row_bg = HEX_IVORY if idx % 2 == 0 else HEX_WHITE
        
        for cell in row_cells:
            set_cell_background(cell, row_bg)
            set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            
        row_cells[0].width = col_widths[0]
        row_cells[1].width = col_widths[1]
        
        # Commit hash column
        p_sha = row_cells[0].paragraphs[0]
        p_sha.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sha = p_sha.add_run(sha)
        run_sha.font.name = "Arial"
        run_sha.font.size = Pt(9)
        run_sha.font.bold = True
        run_sha.font.color.rgb = COLOR_BRASS
        
        # Log message column
        p_msg = row_cells[1].paragraphs[0]
        p_msg.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_msg.paragraph_format.line_spacing = 1.15
        run_msg = p_msg.add_run(msg)
        run_msg.font.name = "Arial"
        run_msg.font.size = Pt(9.5)
        run_msg.font.color.rgb = COLOR_CHARCOAL
        
        # Row layout constraint (keep together)
        trPr_row = table.rows[idx + 1]._tr.get_or_add_trPr()
        trPr_row.append(OxmlElement('w:cantSplit'))
        
    # Number pages in footer
    footer_para = doc.sections[0].footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_foot = footer_para.add_run("Benjamin Barker Studios  ·  Changelog  ·  Page ")
    run_foot.font.name = "Arial"
    run_foot.font.size = Pt(8.5)
    run_foot.font.color.rgb = COLOR_CHARCOAL
    
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    footer_para._p.append(fld)
    
    doc.save(output_name)
    print(f"\nSuccess! Built ledger with {len(commits)} commits.")
    print(f"File saved as: {output_name}\n")

def main():
    print("Reading Git history...")
    raw_log = run_git_log()
    commits = parse_commits(raw_log)
    build_docx(commits)

if __name__ == "__main__":
    main()
